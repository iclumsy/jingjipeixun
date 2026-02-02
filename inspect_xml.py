from docx import Document
from docx.document import Document as _Document
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph
import sys

def iter_block_items(parent):
    if isinstance(parent, _Document):
        parent_elm = parent.element.body
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    else:
        raise ValueError("something's not right")

    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)

doc = Document('training_system/template.docx')

print("--- Searching for Image ---")
# Traverse to find where the image is
for i, table in enumerate(doc.tables):
    for r, row in enumerate(table.rows):
        for c, cell in enumerate(row.cells):
            # Check for drawing/blip
            xml = cell._tc.xml
            if 'w:drawing' in xml:
                print(f"Found drawing in Table {i}, Row {r}, Cell {c}")
                # Try to extract blip ID
                if 'r:embed="' in xml:
                    start = xml.find('r:embed="') + 9
                    end = xml.find('"', start)
                    print(f"  Blip ID: {xml[start:end]}")
                print(f"  Text in cell: {cell.text[:20]}...")

print("\n--- Inspecting Text Formatting ---")
# Check the "Name" value cell
# From previous run: "姓　名" is likely in T0, R0 or similar.
# Let's find "杨洋"
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            if "杨洋" in cell.text:
                print(f"Found '杨洋' in cell. XML fragment:")
                # Print run properties to see font
                for p in cell.paragraphs:
                    for run in p.runs:
                        if "杨洋" in run.text:
                            print(f"  Run text: {run.text}")
                            print(f"  Run XML: {run._element.xml}")
