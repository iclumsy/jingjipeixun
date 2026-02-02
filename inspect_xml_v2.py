from docx import Document
from docx.document import Document as _Document
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph
import sys

doc = Document('training_system/template.docx')

print("--- Searching for Image in Paragraphs ---")
for i, p in enumerate(doc.paragraphs):
    xml = p._element.xml
    if 'w:drawing' in xml:
        print(f"Found w:drawing in Paragraph {i}")
        if 'r:embed="' in xml:
             start = xml.find('r:embed="') + 9
             end = xml.find('"', start)
             print(f"  Blip ID: {xml[start:end]}")
    if 'w:pict' in xml:
        print(f"Found w:pict in Paragraph {i}")
        if 'r:id="' in xml:
             start = xml.find('r:id="') + 6
             end = xml.find('"', start)
             print(f"  Rel ID: {xml[start:end]}")

print("\n--- Searching for Image in Tables (Deep Search) ---")
for i, table in enumerate(doc.tables):
    for r, row in enumerate(table.rows):
        for c, cell in enumerate(row.cells):
            # Check all paragraphs in cell
            for p in cell.paragraphs:
                xml = p._element.xml
                if 'w:drawing' in xml:
                    print(f"Found w:drawing in Table {i}, Row {r}, Cell {c}")
                    # Extract all blip IDs
                    import re
                    ids = re.findall(r'r:embed="([^"]+)"', xml)
                    print(f"  Blip IDs: {ids}")
                if 'w:pict' in xml:
                    print(f"Found w:pict in Table {i}, Row {r}, Cell {c}")
                    ids = re.findall(r'r:id="([^"]+)"', xml)
                    print(f"  Rel IDs: {ids}")
                    # Also check v:imagedata
                    ids2 = re.findall(r'v:imagedata[^>]*r:id="([^"]+)"', xml)
                    print(f"  Imagedata IDs: {ids2}")

