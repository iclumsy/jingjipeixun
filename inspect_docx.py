from docx import Document
import sys

try:
    doc = Document('training_system/template.docx')
    for para in doc.paragraphs:
        if para.text.strip():
            print(f"P: {para.text}")
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text for cell in row.cells]
            print(f"T: {' | '.join(row_text)}")
except Exception as e:
    print(e)
