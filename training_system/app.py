import os
import sqlite3
import shutil
from datetime import datetime
from flask import Flask, render_template, request, jsonify, send_from_directory
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from PIL import Image, ImageOps, ImageStat
import numpy as np
try:
    import cv2
except Exception:
    cv2 = None
from lxml import etree
import io
from werkzeug.utils import secure_filename
import re

app = Flask(__name__)
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
app.config['UPLOAD_FOLDER'] = os.path.join(BASE_DIR, 'uploads')
app.config['REVIEWED_FOLDER'] = os.path.join(BASE_DIR, 'reviewed_students')
app.config['DATABASE'] = os.path.join(BASE_DIR, 'database/students.db')
app.config['TEMPLATE_PATH'] = os.path.join(BASE_DIR, '特种设备作业人员考试体检表（锅炉水处理、客运索道司机）-杜臻.docx')
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB max upload

# Ensure directories exist
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
os.makedirs(app.config['REVIEWED_FOLDER'], exist_ok=True)
os.makedirs(os.path.join(BASE_DIR, 'database'), exist_ok=True)

def get_db_connection():
    conn = sqlite3.connect(app.config['DATABASE'])
    conn.row_factory = sqlite3.Row
    return conn

def init_db():
    conn = get_db_connection()
    conn.execute('''
        CREATE TABLE IF NOT EXISTS students (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            name TEXT NOT NULL,
            gender TEXT NOT NULL,
            education TEXT NOT NULL,
            school TEXT,
            major TEXT,
            id_card TEXT NOT NULL,
            phone TEXT NOT NULL,
            company TEXT,
            company_address TEXT,
            job_category TEXT NOT NULL,
            exam_project TEXT,
            exam_code TEXT,
            exam_category TEXT NOT NULL,
            status TEXT DEFAULT 'unreviewed',
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            photo_path TEXT,
            diploma_path TEXT,
            cert_path TEXT,
            id_card_front_path TEXT,
            id_card_back_path TEXT,
            training_form_path TEXT,
            theory_exam_time TEXT,
            practical_exam_time TEXT,
            passed TEXT,
            theory_makeup_time TEXT,
            makeup_exam TEXT
        )
    ''')
    conn.commit()
    conn.close()

init_db()

# Mapping for human-readable suffixes used when saving files
LABEL_NAME_MAP = {
    'photo': '个人照片',
    'diploma': '学历证书',
    'cert': '所持证件',
    'id_card_front': '身份证正面',
    'id_card_back': '身份证反面'
}


def _ensure_ext(fname, default_ext='.jpg'):
    _, ext = os.path.splitext(fname)
    return ext.lower() if ext else default_ext


def process_and_save_file(file_storage, id_card, label_key):
    """Save uploaded file using pattern '<id_card>-<label>.<ext>'.
    If label_key is 'photo', run face-detection centering and produce a one-inch image.
    Returns relative path like 'uploads/....'
    """
    if not file_storage or not file_storage.filename:
        return ''

    label_name = LABEL_NAME_MAP.get(label_key, label_key)
    orig_ext = _ensure_ext(file_storage.filename, '.jpg')
    # Do NOT use secure_filename on entire name as it strips Chinese characters
    safe_name = f"{id_card}-{label_name}{orig_ext}"
    abs_path = os.path.join(app.config['UPLOAD_FOLDER'], safe_name)

    # Save initial upload (overwrite if exists)
    file_storage.save(abs_path)

    # If this is the personal photo, perform face-detection centering and one-inch generation
    if label_key == 'photo':
        try:
            _process_photo_center_face(abs_path)
        except Exception:
            # If processing fails, leave the original upload
            pass

    return f"uploads/{safe_name}"


def _process_photo_center_face(abs_path):
    """Load image at abs_path, detect face with OpenCV (if available),
    generate a one-inch (portrait 2.5x3.5cm) image where the face is centered.
    Overwrites file at abs_path with a JPEG.
    """
    # Target physical size: 2.5cm x 3.5cm -> inches
    tgt_w_in = 2.5 / 2.54
    tgt_h_in = 3.5 / 2.54
    dpi = 300
    px_w = max(120, int(tgt_w_in * dpi))
    px_h = max(160, int(tgt_h_in * dpi))

    im = Image.open(abs_path).convert('RGBA')
    # flatten alpha onto white
    white_bg = Image.new('RGBA', im.size, (255, 255, 255, 255))
    composed = Image.alpha_composite(white_bg, im).convert('RGB')
    img_w, img_h = composed.size

    # Try face detection with OpenCV
    face_center = None
    face_box = None
    if cv2 is not None:
        try:
            gray = cv2.cvtColor(np.array(composed), cv2.COLOR_RGB2GRAY)
            cascade_path = cv2.data.haarcascades + 'haarcascade_frontalface_default.xml'
            face_cascade = cv2.CascadeClassifier(cascade_path)
            faces = face_cascade.detectMultiScale(gray, scaleFactor=1.1, minNeighbors=5, minSize=(30, 30))
            if len(faces) > 0:
                # choose largest face
                faces = sorted(faces, key=lambda r: r[2] * r[3], reverse=True)
                x, y, w, h = faces[0]
                face_center = (x + w / 2.0, y + h / 2.0)
                face_box = (x, y, w, h)
        except Exception:
            face_center = None

    # If we detected a face, perform cover-scaling and center the face
    if face_center is not None:
        try:
            # cover scale so that final crop will be filled (may crop edges)
            scale = max(px_w / img_w, px_h / img_h)
            resized_w = max(1, int(img_w * scale))
            resized_h = max(1, int(img_h * scale))
            resized = composed.resize((resized_w, resized_h), Image.LANCZOS)

            # scaled face center
            fc_x = int(round(face_center[0] * scale))
            fc_y = int(round(face_center[1] * scale))

            # compute top-left corner on resized image so face_center maps to canvas center
            src_left = int(fc_x - px_w // 2)
            src_top = int(fc_y - px_h // 2)
            # clamp
            src_left = max(0, min(resized_w - px_w, src_left))
            src_top = max(0, min(resized_h - px_h, src_top))

            cropped = resized.crop((src_left, src_top, src_left + px_w, src_top + px_h))
            out = Image.new('RGB', (px_w, px_h), (255, 255, 255))
            out.paste(cropped, (0, 0))
            out.save(abs_path, format='JPEG', quality=95)
            return
        except Exception:
            pass

    # Fallback: contain-scale and center (no face centering)
    try:
        scale = min(px_w / img_w, px_h / img_h)
        new_w = max(1, int(img_w * scale))
        new_h = max(1, int(img_h * scale))
        resized = composed.resize((new_w, new_h), Image.LANCZOS)
        canvas = Image.new('RGB', (px_w, px_h), (255, 255, 255))
        paste_left = (px_w - new_w) // 2
        paste_top = (px_h - new_h) // 2
        canvas.paste(resized, (paste_left, paste_top))
        canvas.save(abs_path, format='JPEG', quality=95)
    except Exception:
        # last resort: save original flattened as JPEG
        composed.convert('RGB').save(abs_path, format='JPEG', quality=90)

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/admin')
def admin():
    return render_template('admin.html')

@app.route('/api/students', methods=['POST'])
def create_student():
    try:
        data = request.form
        files = request.files
        
        # Server-side validation (required fields and patterns)
        required_fields = ['name', 'gender', 'education', 'id_card', 'phone', 'job_category', 'exam_category']
        errors = {}
        for f in required_fields:
            if not data.get(f):
                errors[f] = '必填项'
        if 'gender' in data and data.get('gender') not in ['男', '女']:
            errors['gender'] = '性别须为“男”或“女”'
        if 'id_card' in data and not re.fullmatch(r'\d{17}[\dXx]', data.get('id_card', '')):
            errors['id_card'] = '身份证号格式不正确'
        if 'phone' in data and not re.fullmatch(r'\d{11}', data.get('phone', '')):
            errors['phone'] = '手机号格式不正确'
        if errors:
            return jsonify({'error': 'validation_failed', 'fields': errors}), 400
        
        # Save files using ID card + descriptive name convention
        file_paths = {}
        # Map input names to DB column keys
        file_map = {
            'photo': 'photo_path',
            'diploma': 'diploma_path',
            'cert': 'cert_path',
            'id_card_front': 'id_card_front_path',
            'id_card_back': 'id_card_back_path'
        }

        id_card_val = data.get('id_card', '').strip()
        for input_name, db_key in file_map.items():
            file = files.get(input_name)
            if file and file.filename and id_card_val:
                try:
                    rel = process_and_save_file(file, id_card_val, input_name)
                    file_paths[db_key] = rel
                except Exception:
                    file_paths[db_key] = ""
            else:
                file_paths[db_key] = ""
        
        # training_form is generated, not collected
        file_paths['training_form_path'] = ""

        conn = get_db_connection()
        cursor = conn.cursor()
        cursor.execute('''
            INSERT INTO students (
                name, gender, education, school, major, id_card, phone,
                company, company_address, job_category, exam_project, exam_code,
                exam_category, photo_path, diploma_path, cert_path, id_card_front_path, id_card_back_path, training_form_path,
                theory_exam_time, practical_exam_time, passed, theory_makeup_time, makeup_exam
            ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        ''', (
            data['name'], data['gender'], data['education'], data.get('school', ''), data.get('major', ''),
            data['id_card'], data['phone'], data.get('company', ''), data.get('company_address', ''),
            data['job_category'], data.get('exam_project', ''), data.get('exam_code', ''), data['exam_category'],
            file_paths['photo_path'], file_paths['diploma_path'], file_paths['cert_path'],
            file_paths['id_card_front_path'], file_paths['id_card_back_path'], file_paths['training_form_path'],
            data.get('theory_exam_time', ''), data.get('practical_exam_time', ''), data.get('passed', ''),
            data.get('theory_makeup_time', ''), data.get('makeup_exam', '')
        ))
        conn.commit()
        conn.close()
        return jsonify({'message': 'Student added successfully'}), 201
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/students', methods=['GET'])
def get_students():
    status = request.args.get('status', 'unreviewed')
    search = request.args.get('search', '')
    company = request.args.get('company', '')
    passed = request.args.get('passed', '')
    examined = request.args.get('examined', '')
    
    conn = get_db_connection()
    
    # Base query
    if status == 'examined':
        # Special case for examined status
        query = "SELECT * FROM students WHERE ((theory_exam_time IS NOT NULL AND theory_exam_time != '') OR (practical_exam_time IS NOT NULL AND practical_exam_time != ''))"
        params = []
    else:
        query = "SELECT * FROM students WHERE status = ?"
        params = [status]
    
    if search:
        query += " AND (name LIKE ? OR id_card LIKE ? OR phone LIKE ?)"
        params.extend([f"%{search}%", f"%{search}%", f"%{search}%"])
    
    if company:
        query += " AND company LIKE ?"
        params.append(f"%{company}%")
    
    if passed:
        query += " AND passed = ?"
        params.append(passed)
    
    if examined and status != 'examined':
        query += " AND ((theory_exam_time IS NOT NULL AND theory_exam_time != '') OR (practical_exam_time IS NOT NULL AND practical_exam_time != ''))"
        
    students = conn.execute(query, params).fetchall()
    conn.close()
    
    return jsonify([dict(s) for s in students])

@app.route('/api/students/<int:id>', methods=['PUT', 'PATCH'])
def update_student(id):
    allowed_text = [
        'name', 'gender', 'education', 'school', 'major', 'id_card', 'phone',
        'company', 'company_address', 'job_category', 'exam_project', 'exam_code',
        'exam_category', 'theory_exam_time', 'practical_exam_time', 'passed',
        'theory_makeup_time', 'makeup_exam'
    ]
    file_map = {
        'photo': 'photo_path',
        'diploma': 'diploma_path',
        'cert': 'cert_path',
        'id_card_front': 'id_card_front_path',
        'id_card_back': 'id_card_back_path'
    }
    conn = get_db_connection()
    current = conn.execute('SELECT * FROM students WHERE id = ?', (id,)).fetchone()
    if not current:
        conn.close()
        return jsonify({'error': 'Student not found'}), 404
    updates = {}
    # Prefer form data when present (multipart), else JSON
    if request.form:
        data = request.form
        for k in allowed_text:
            if k in data:
                updates[k] = data[k]
        # Validate present text fields (partial update)
        errors = {}
        if 'gender' in data and data.get('gender') not in ['男', '女']:
            errors['gender'] = '性别须为“男”或“女”'
        if 'id_card' in data and not re.fullmatch(r'\d{17}[\dXx]', data.get('id_card', '')):
            errors['id_card'] = '身份证号格式不正确'
        if 'phone' in data and not re.fullmatch(r'\d{11}', data.get('phone', '')):
            errors['phone'] = '手机号格式不正确'
        if errors:
            conn.close()
            return jsonify({'error': 'validation_failed', 'fields': errors}), 400
        for input_name, db_key in file_map.items():
            f = request.files.get(input_name)
            if f and f.filename:
                    # Determine id_card to use for naming (prefer updated one if provided)
                    id_card_for_name = data.get('id_card', current['id_card'])
                    # delete old file if exists
                    old_rel = current[db_key]
                    if old_rel:
                        old_fn = os.path.basename(old_rel)
                        old_abs = os.path.join(app.config['UPLOAD_FOLDER'], old_fn)
                        if os.path.exists(old_abs):
                            try:
                                os.remove(old_abs)
                            except Exception:
                                pass
                    try:
                        rel = process_and_save_file(f, id_card_for_name, input_name)
                        updates[db_key] = rel
                    except Exception:
                        updates[db_key] = ''
    else:
        payload = request.get_json(silent=True) or {}
        for k in allowed_text:
            if k in payload:
                updates[k] = payload[k]
        # Validate present text fields (partial update)
        errors = {}
        if 'gender' in payload and payload.get('gender') not in ['男', '女']:
            errors['gender'] = '性别须为“男”或“女”'
        if 'id_card' in payload and not re.fullmatch(r'\d{17}[\dXx]', payload.get('id_card', '')):
            errors['id_card'] = '身份证号格式不正确'
        if 'phone' in payload and not re.fullmatch(r'\d{11}', payload.get('phone', '')):
            errors['phone'] = '手机号格式不正确'
        if errors:
            conn.close()
            return jsonify({'error': 'validation_failed', 'fields': errors}), 400
    if not updates:
        conn.close()
        return jsonify({'error': 'no fields to update'}), 400
    sets = ', '.join([f"{k} = ?" for k in updates.keys()])
    values = list(updates.values()) + [id]
    conn.execute(f"UPDATE students SET {sets} WHERE id = ?", values)
    conn.commit()
    student = conn.execute('SELECT * FROM students WHERE id = ?', (id,)).fetchone()
    conn.close()
    return jsonify(dict(student))

@app.route('/api/students/<int:id>/reject', methods=['POST'])
def reject_student(id):
    conn = get_db_connection()
    student = conn.execute('SELECT * FROM students WHERE id = ?', (id,)).fetchone()
    
    if not student:
        conn.close()
        return jsonify({'error': 'Student not found'}), 404
        
    # Delete files
    for key in ['photo_path', 'diploma_path', 'cert_path', 'id_card_front_path', 'id_card_back_path', 'training_form_path']:
        if student[key]:
            # Construct absolute path from relative DB path
            # DB stores "uploads/filename", we need "/abs/path/to/uploads/filename"
            # But UPLOAD_FOLDER is "/abs/path/to/uploads"
            # So we strip "uploads/" from student[key] or use basename
            filename = os.path.basename(student[key])
            file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            if os.path.exists(file_path):
                os.remove(file_path)
            
    conn.execute('DELETE FROM students WHERE id = ?', (id,))
    conn.commit()
    conn.close()
    return jsonify({'message': 'Student rejected and deleted'})

@app.route('/api/students/<int:id>/approve', methods=['POST'])
def approve_student(id):
    conn = get_db_connection()
    exists = conn.execute('SELECT id FROM students WHERE id = ?', (id,)).fetchone()
    if not exists:
        conn.close()
        return jsonify({'error': 'Student not found'}), 404
    conn.execute("UPDATE students SET status = 'reviewed' WHERE id = ?", (id,))
    conn.commit()
    student = conn.execute('SELECT * FROM students WHERE id = ?', (id,)).fetchone()
    conn.close()
    return jsonify(dict(student))

@app.route('/api/students/<int:id>/generate', methods=['POST'])
def generate_materials(id):
    conn = get_db_connection()
    student = conn.execute('SELECT * FROM students WHERE id = ?', (id,)).fetchone()
    if not student:
        conn.close()
        return jsonify({'error': 'Student not found'}), 404
    # Use only ID card number for reviewed folder name
    folder_name = f"{student['id_card']}"
    student_dir = os.path.join(app.config['REVIEWED_FOLDER'], folder_name)
    os.makedirs(student_dir, exist_ok=True)
    file_mapping = {
        'training_form_path': f"{student['id_card']}-培训信息登记表.jpg",
        'diploma_path': f"{student['id_card']}-学历证书复印件.jpg",
        'cert_path': f"{student['id_card']}-所持证件复印件.jpg",
        'id_card_front_path': f"{student['id_card']}-身份证正面.jpg",
        'id_card_back_path': f"{student['id_card']}-身份证反面.jpg",
        'photo_path': f"{student['id_card']}-个人照片.jpg"
    }
    for db_field, target_name in file_mapping.items():
        db_path = student[db_field]
        if db_path:
            filename = os.path.basename(db_path)
            src_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            if os.path.exists(src_path):
                _, ext = os.path.splitext(src_path)
                base_target_name = os.path.splitext(target_name)[0]
                final_target_name = base_target_name + ext
                shutil.copy2(src_path, os.path.join(student_dir, final_target_name))
    doc_path = os.path.join(student_dir, f"{student['id_card']}-体检表.docx")
    photo_abs_path = None
    if student['photo_path']:
        filename = os.path.basename(student['photo_path'])
        photo_abs_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
    generate_word_doc(app.config['TEMPLATE_PATH'], doc_path, student, photo_abs_path)
    # store generated docx relative path in DB and return download url
    rel_path = f"reviewed_students/{folder_name}/{os.path.basename(doc_path)}"
    conn.execute('UPDATE students SET training_form_path = ? WHERE id = ?', (rel_path, id))
    conn.commit()
    conn.close()
    download_url = f"/reviewed_students/{folder_name}/{os.path.basename(doc_path)}"
    return jsonify({'message': 'materials generated', 'download_url': download_url, 'path': rel_path})

def generate_word_doc(template_path, output_path, data, photo_path=None):
    doc = Document(template_path)
    
    # 1. Text Replacement (Preserving Style)
    replacements = {
        '姓　名': data['name'],
        '姓名': data['name'],
        '性别': data['gender'],
        '文化程度': data['education'],
        '身份证号': data['id_card'],
        '移动电话': data['phone'],
        '手机号': data['phone'],
        '工作单位': data['company'],
        '作业类别': data['job_category'],
        '操作项目': data['exam_project']
    }
    
    for table in doc.tables:
        for row in table.rows:
            for i, cell in enumerate(row.cells):
                text = cell.text.strip()
                if text in replacements:
                    # Look for the next cell
                    if i + 1 < len(row.cells):
                        target_cell = row.cells[i+1]
                        
                        # Safety check: Don't overwrite if the target cell is also a known label
                        # This prevents issues with merged cells or repeated labels like "移动电话 | 移动电话"
                        if target_cell.text.strip() in replacements:
                            continue
                            
                        new_val = replacements[text]
                        
                        # Replace text in the first run of the first paragraph to preserve formatting
                        if target_cell.paragraphs:
                            p = target_cell.paragraphs[0]
                            if p.runs:
                                p.runs[0].text = new_val
                                # Clear other runs in this paragraph
                                for r in p.runs[1:]:
                                    r.text = ''
                            else:
                                # No runs, add one (formatting might be default, but better than nothing)
                                p.add_run(new_val)
                                
                            # Clear other paragraphs in the cell
                            for p in target_cell.paragraphs[1:]:
                                p.clear()
                        else:
                            target_cell.text = new_val
    
    # 2. Image Replacement: find image relationship in document and replace blob with processed photo
    if photo_path and os.path.exists(photo_path):
        try:
            # Prepare processed photo: convert to white background, crop center to square and resize to 1 inch (~300px)
            try:
                im = Image.open(photo_path)
            except Exception:
                im = Image.open(photo_path)

            # Ensure RGBA for alpha handling, then flatten over white
            if im.mode != 'RGBA':
                im_rgba = im.convert('RGBA')
            else:
                im_rgba = im
            white_bg = Image.new('RGBA', im_rgba.size, (255, 255, 255, 255))
            composed = Image.alpha_composite(white_bg, im_rgba)
            composed = composed.convert('RGB')

            # Try GrabCut foreground extraction (preferred). Fall back to pixel-threshold replace.
            def grabcut_white_background(pil_img):
                if cv2 is None:
                    return None
                try:
                    img_cv = cv2.cvtColor(np.array(pil_img), cv2.COLOR_RGB2BGR)
                    h, w = img_cv.shape[:2]
                    mask = np.zeros((h, w), np.uint8)
                    rect = (int(w*0.05), int(h*0.05), int(w*0.9), int(h*0.9))
                    bgdModel = np.zeros((1,65), np.float64)
                    fgdModel = np.zeros((1,65), np.float64)
                    cv2.grabCut(img_cv, mask, rect, bgdModel, fgdModel, 5, cv2.GC_INIT_WITH_RECT)
                    # Create mask where sure or likely foreground are 1
                    mask2 = np.where((mask==cv2.GC_FGD) | (mask==cv2.GC_PR_FGD), 255, 0).astype('uint8')
                    # Smooth mask and apply alpha
                    mask_blur = cv2.GaussianBlur(mask2, (5,5), 0)
                    alpha = (mask_blur/255.0).astype('float32')
                    b,g,r = cv2.split(img_cv)
                    img_rgb = cv2.merge([r,g,b]).astype('float32')
                    white = np.ones_like(img_rgb) * 255.0
                    out = img_rgb * alpha[:,:,None] + white * (1 - alpha[:,:,None])
                    out = np.clip(out, 0, 255).astype('uint8')
                    return Image.fromarray(out)
                except Exception:
                    return None

            replaced_img = None
            try:
                replaced_img = grabcut_white_background(composed)
            except Exception:
                replaced_img = None

            if replaced_img is None:
                # Fallback to pixel-threshold replacement
                def replace_bg_with_white(img_obj, bg_color, thr=40):
                    img2 = img_obj.convert('RGBA')
                    px2 = img2.load()
                    ww, hh = img2.size
                    for yy in range(hh):
                        for xx in range(ww):
                            r, g, b, a = px2[xx, yy]
                            if a == 0:
                                px2[xx, yy] = (255, 255, 255, 255)
                                continue
                            if (abs(r - bg_color[0]) + abs(g - bg_color[1]) + abs(b - bg_color[2])) <= thr:
                                px2[xx, yy] = (255, 255, 255, 255)
                    return img2.convert('RGB')

                # estimate background color from corners
                w, h = composed.size
                cs_w = max(1, w // 20)
                cs_h = max(1, h // 20)
                corner_regions = []
                for sx, sy in [(0, 0), (w - cs_w, 0), (0, h - cs_h), (w - cs_w, h - cs_h)]:
                    try:
                        region = composed.crop((sx, sy, sx + cs_w, sy + cs_h))
                        stat = ImageStat.Stat(region)
                        corner_regions.append(tuple(int(c) for c in stat.mean))
                    except Exception:
                        corner_regions.append((255, 255, 255))

                try:
                    bg_color = max(set(corner_regions), key=corner_regions.count)
                except Exception:
                    bg_color = corner_regions[0]

                composed = replace_bg_with_white(composed, bg_color, thr=50)
            else:
                composed = replaced_img

            # Do not perform additional cropping here — uploaded photo is already
            # prepared (face-centered one-inch) on save. Use the composed image as-is.
            default_w_in = 2.5 / 2.54  # ~0.984in
            default_h_in = 3.5 / 2.54  # ~1.378in
            bio_tmp = io.BytesIO()
            composed.save(bio_tmp, format='JPEG', quality=95)
            src_bytes = bio_tmp.getvalue()

            # First: try inserting into a table cell (preferred)
            replaced = False
            target_cell = None
            found_table = None
            found_row_idx = None
            found_col_idx = None
            # First look for a cell that contains the word '照' or '照片' (common markers)
            for ti, table in enumerate(doc.tables):
                for ri, row in enumerate(table.rows):
                    for ci, cell in enumerate(row.cells):
                        txt = cell.text.strip()
                        if '照片' in txt or '照' in txt:
                            target_cell = cell
                            found_table = table
                            found_row_idx = ri
                            found_col_idx = ci
                            break
                    if target_cell:
                        break
                if target_cell:
                    break

            # Fallback: choose top-right cell of the first table
            if not target_cell and len(doc.tables) > 0:
                tbl = doc.tables[0]
                target_cell = tbl.cell(0, len(tbl.rows[0].cells)-1)
                found_table = tbl
                found_row_idx = 0
                found_col_idx = len(tbl.rows[0].cells)-1

            if target_cell:
                # Determine target cell size in inches
                ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
                try:
                    cell_root = etree.fromstring(target_cell._tc.xml.encode('utf-8'))
                except Exception:
                    cell_root = None
                tcW = None
                if cell_root is not None:
                    tcW = cell_root.xpath('.//w:tcPr/w:tcW', namespaces=ns)
                if tcW:
                    try:
                        w_val = int(tcW[0].get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}w'))
                        w_type = tcW[0].get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}type')
                        if w_type == 'dxa':
                            target_w_in = w_val / 1440.0
                        else:
                            target_w_in = default_w_in
                    except Exception:
                        target_w_in = default_w_in
                else:
                    # fallback to table gridCol width
                    try:
                        tbl_root = etree.fromstring(found_table._tbl.xml.encode('utf-8'))
                        gridCols = tbl_root.xpath('.//w:tblGrid/w:gridCol', namespaces=ns)
                        if gridCols and found_col_idx is not None and found_col_idx < len(gridCols):
                            gv = int(gridCols[found_col_idx].get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}w'))
                            target_w_in = gv / 1440.0
                        else:
                            target_w_in = default_w_in
                    except Exception:
                        target_w_in = default_w_in

                # We'll derive target height from target width to avoid tiny template row-height values
                try:
                    if 'target_w_in' in locals():
                        if not target_w_in or target_w_in <= 0 or target_w_in > 20:
                            target_w_in = None
                except Exception:
                    target_w_in = None

                if 'target_w_in' in locals() and target_w_in:
                    tgt_w = target_w_in
                    # assume portrait 2.5x3.5cm aspect ratio (one-inch portrait)
                    tgt_h = tgt_w * (3.5 / 2.5)
                else:
                    tgt_w = default_w_in
                    tgt_h = default_h_in

                # Ensure minimum sensible physical size to avoid extremely thin image
                min_w_in = 0.6
                min_h_in = 0.8
                if tgt_w < min_w_in:
                    tgt_w = min_w_in
                if tgt_h < min_h_in:
                    tgt_h = min_h_in

                # Crop source image to target aspect ratio (cover)
                try:
                    # Load source image (already background-processed in src_bytes)
                    src_img = Image.open(io.BytesIO(src_bytes)).convert('RGB')
                    src_w, src_h = src_img.size

                    # resize to fit target box while preserving original aspect ratio
                    dpi = 300
                    px_w = max(120, int(tgt_w * dpi))
                    px_h = max(160, int(tgt_h * dpi))
                    scale = min(px_w / src_w, px_h / src_h)
                    new_w = max(1, int(src_w * scale))
                    new_h = max(1, int(src_h * scale))
                    resized_img = src_img.resize((new_w, new_h), Image.LANCZOS)

                    # Paste onto white canvas of exact target pixels so we don't over-crop
                    canvas = Image.new('RGB', (px_w, px_h), (255, 255, 255))
                    paste_left = (px_w - new_w) // 2
                    paste_top = (px_h - new_h) // 2
                    canvas.paste(resized_img, (paste_left, paste_top))

                    out_bio = io.BytesIO()
                    canvas.save(out_bio, format='JPEG', quality=95)
                    img_bytes_final = out_bio.getvalue()

                    # clear existing paragraphs
                    for p in list(target_cell.paragraphs):
                        p.clear()
                    p = target_cell.paragraphs[0]
                    try:
                        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    except Exception:
                        pass
                    run = p.add_run()
                    bio2 = io.BytesIO(img_bytes_final)
                    # insert with target sizes
                    try:
                        run.add_picture(bio2, width=Inches(tgt_w), height=Inches(tgt_h))
                    except TypeError:
                        bio2.seek(0)
                        run.add_picture(bio2, width=Inches(tgt_w))
                    # make sure fallback replacement uses this final bytes
                    img_bytes = img_bytes_final
                    replaced = True
                except Exception as e:
                    print(f"Error preparing/resizing picture for cell: {e}")

            # If insertion didn't happen, fall back to replacing existing image rels
            if not replaced:
                # Find first image relationship (blip) in document part relationships
                blip_rid = None
                for rel in doc.part.rels:
                    # relationship objects have reltype; image reltypes contain 'image'
                    try:
                        rel_obj = doc.part.rels[rel]
                        if getattr(rel_obj, 'reltype', '').endswith('/image') or 'image' in getattr(rel_obj.target_part, '__class__', '').lower():
                            blip_rid = rel
                            break
                    except Exception:
                        continue

                # fallback: search paragraphs xml for r:embed
                if not blip_rid:
                    import re
                    for p in doc.paragraphs:
                        xml = p._element.xml
                        if 'r:embed="' in xml:
                            match = re.search(r'r:embed="([^\"]+)"', xml)
                            if match:
                                blip_rid = match.group(1)
                                break

                if blip_rid and blip_rid in doc.part.rels:
                    try:
                        doc.part.rels[blip_rid].target_part._blob = img_bytes
                        replaced = True
                    except Exception:
                        try:
                            with open(output_path + '.tmp.jpg', 'wb') as f:
                                f.write(img_bytes)
                            doc.part.rels[blip_rid].target_part._blob = img_bytes
                            replaced = True
                        except Exception:
                            replaced = False
        except Exception as e:
            print(f"Error processing or replacing image: {e}")

    doc.save(output_path)

@app.route('/uploads/<path:filename>')
def serve_uploads(filename):
    return send_from_directory(app.config['UPLOAD_FOLDER'], filename)

@app.route('/reviewed_students/<path:filename>')
def serve_reviewed(filename):
    return send_from_directory(app.config['REVIEWED_FOLDER'], filename)

@app.route('/api/companies', methods=['GET'])
def get_companies():
    conn = get_db_connection()
    companies = conn.execute('SELECT DISTINCT company FROM students WHERE company IS NOT NULL AND company != "" ORDER BY company').fetchall()
    conn.close()
    return jsonify([dict(c)['company'] for c in companies])

@app.route('/api/students/batch/approve', methods=['POST'])
def batch_approve_students():
    try:
        data = request.get_json()
        if not data or 'ids' not in data:
            return jsonify({'error': 'Missing student IDs'}), 400
        
        ids = data['ids']
        if not isinstance(ids, list):
            return jsonify({'error': 'IDs must be a list'}), 400
        
        conn = get_db_connection()
        
        # Update status for all selected students
        placeholders = ','.join(['?'] * len(ids))
        query = f"UPDATE students SET status = 'reviewed' WHERE id IN ({placeholders})"
        conn.execute(query, ids)
        conn.commit()
        conn.close()
        
        return jsonify({'message': f'Successfully approved {len(ids)} students'}), 200
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/students/batch/reject', methods=['POST'])
def batch_reject_students():
    try:
        data = request.get_json()
        if not data or 'ids' not in data:
            return jsonify({'error': 'Missing student IDs'}), 400
        
        ids = data['ids']
        if not isinstance(ids, list):
            return jsonify({'error': 'IDs must be a list'}), 400
        
        conn = get_db_connection()
        
        # Get all students to delete
        placeholders = ','.join(['?'] * len(ids))
        students = conn.execute(f"SELECT * FROM students WHERE id IN ({placeholders})", ids).fetchall()
        
        # Delete files for each student
        for student in students:
            for key in ['photo_path', 'diploma_path', 'cert_path', 'id_card_front_path', 'id_card_back_path', 'training_form_path']:
                if student[key]:
                    filename = os.path.basename(student[key])
                    file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
                    if os.path.exists(file_path):
                        os.remove(file_path)
        
        # Delete students from database
        conn.execute(f"DELETE FROM students WHERE id IN ({placeholders})", ids)
        conn.commit()
        conn.close()
        
        return jsonify({'message': f'Successfully rejected and deleted {len(ids)} students'}), 200
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/students/batch/delete', methods=['POST'])
def batch_delete_students():
    try:
        data = request.get_json()
        if not data or 'ids' not in data:
            return jsonify({'error': 'Missing student IDs'}), 400
        
        ids = data['ids']
        if not isinstance(ids, list):
            return jsonify({'error': 'IDs must be a list'}), 400
        
        conn = get_db_connection()
        
        # Get all students to delete
        placeholders = ','.join(['?'] * len(ids))
        students = conn.execute(f"SELECT * FROM students WHERE id IN ({placeholders})", ids).fetchall()
        
        # Delete files for each student
        for student in students:
            for key in ['photo_path', 'diploma_path', 'cert_path', 'id_card_front_path', 'id_card_back_path', 'training_form_path']:
                if student[key]:
                    filename = os.path.basename(student[key])
                    file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
                    if os.path.exists(file_path):
                        os.remove(file_path)
        
        # Delete students from database
        conn.execute(f"DELETE FROM students WHERE id IN ({placeholders})", ids)
        conn.commit()
        conn.close()
        
        return jsonify({'message': f'Successfully deleted {len(ids)} students'}), 200
    except Exception as e:
        return jsonify({'error': str(e)}), 500

if __name__ == '__main__':
    app.run(debug=True, port=5000)
