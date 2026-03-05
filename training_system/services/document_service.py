"""
文档生成服务。

本模块负责根据学员数据和 Word 模板自动生成体检表文档。

核心功能:
    - 判断学员的操作项目是否需要体检表
    - 从 .docx 模板中查找表格单元格并填入学员数据
    - 将学员证件照处理后插入模板中的照片区域

支持的体检表类型:
    - 叉车司机体检表 (项目代号 N1)
    - 锅炉水处理体检表 (项目代号 G3)

模板文件存放在项目根目录下，文件名需与 HEALTH_CHECK_TEMPLATES 中定义的一致。

文档填写逻辑:
    模板中使用"标签-值"的表格布局，如:
    | 姓名 |        | 性别 |      |
    标签单元格与值单元格水平相邻，程序通过匹配标签文本
    自动将数据填入右侧的值单元格。
"""
import os
import io
import tempfile
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from PIL import Image
from lxml import etree
from flask import current_app
from services.image_service import change_id_photo_bg


# ======================== 体检表模板配置 ========================

# 操作项目名称 -> 模板文件名映射
# 新增体检表类型时，在此添加映射并将模板文件放入项目根目录
HEALTH_CHECK_TEMPLATES = {
    '叉车司机': '叉车司机体检表.docx',
    '锅炉水处理': '锅炉水处理体检表.docx'
}

# 项目代号 -> 操作项目名称映射
# 优先通过代号匹配，比名称匹配更精确
HEALTH_CHECK_PROJECT_CODES = {
    'N1': '叉车司机',     # N1: 场（厂）内专用机动车辆 - 叉车
    'G3': '锅炉水处理'   # G3: 锅炉压力容器 - 水处理
}


def needs_health_check(exam_project, project_code=''):
    """
    检查该操作项目是否需要生成体检表。
    
    参数:
        exam_project: 操作项目名称
        project_code: 项目代号（如 G3）
        
    返回:
        tuple: (needs_check: bool, template_key: str 或 None)
    """
    # 优先通过项目代号精确匹配（如 N1、G3）
    normalized_code = (project_code or '').strip().upper()
    if normalized_code in HEALTH_CHECK_PROJECT_CODES:
        return True, HEALTH_CHECK_PROJECT_CODES[normalized_code]

    if not exam_project:
        return False, None
    
    # 降级：通过操作项目名称模糊匹配（包含即命中）
    normalized_project = str(exam_project).strip()
    for key in HEALTH_CHECK_TEMPLATES.keys():
        if key in normalized_project:
            return True, key
    return False, None


def generate_health_check_form(student, base_dir, students_folder):
    """
    为符合条件的学员生成体检表。
    
    参数:
        student: 学员数据字典
        base_dir: 基础目录路径
        students_folder: 学员文件夹路径
        
    返回:
        str: 生成文件的相对路径，不需要时返回 None
    """
    # 判断该学员是否需要体检表
    exam_project = student.get('exam_project', '')
    project_code = student.get('project_code', '')
    needs_check, template_key = needs_health_check(exam_project, project_code)
    
    if not needs_check:
        return None
    
    # 查找对应的模板文件
    template_name = HEALTH_CHECK_TEMPLATES.get(template_key)
    if not template_name:
        return None
    
    template_path = os.path.join(base_dir, template_name)
    if not os.path.exists(template_path):
        current_app.logger.warning(f'Health check template not found: {template_path}')
        return None
    
    # 构建学员文件夹路径，格式：students/<培训类型>-<公司>-<姓名>/
    training_type = student.get('training_type', 'special_operation')
    training_type_name = '特种作业' if training_type == 'special_operation' else '特种设备'
    student_folder_name = f"{training_type_name}-{student.get('company', '')}-{student['name']}"
    student_folder_path = os.path.join(students_folder, student_folder_name)
    os.makedirs(student_folder_path, exist_ok=True)
    
    # 生成体检表文件路径，格式：<身份证号>-<姓名>-体检表.docx
    doc_path = os.path.join(
        student_folder_path,
        f"{student['id_card']}-{student['name']}-体检表.docx"
    )
    
    # 获取学员照片的绝对路径（用于插入体检表）
    photo_abs_path = None
    if student.get('photo_path'):
        photo_abs_path = os.path.join(base_dir, student['photo_path'])
    
    # 准备模板填充数据
    data = {
        'name': student['name'],
        'gender': student['gender'],
        'id_card': student['id_card']
    }
    
    # 生成 Word 文档
    generate_word_doc(template_path, doc_path, data, photo_abs_path)
    
    # 返回相对路径（用于数据库存储）
    rel_path = f"students/{student_folder_name}/{os.path.basename(doc_path)}"
    return rel_path


def generate_word_doc(template_path, output_path, data, photo_path=None):
    """
    根据模板生成填写学员数据的 Word 文档。

    参数:
        template_path: 模板文档路径
        output_path: 输出文档保存路径
        data: 包含学员数据的字典
        photo_path: 学员照片路径（可选）
    """
    try:
        doc = Document(template_path)

        # ---- 第一步：文本替换（保留原格式样式） ----
        # 定义标签文本 -> 填充值的映射
        replacements = {
            '姓名': data.get('name', ''),
            '性别': data.get('gender', ''),
            '身份证号': data.get('id_card', ''),
        }

        # 遍历文档中所有表格，查找标签单元格并填入对应的值
        for table in doc.tables:
            for row in table.rows:
                for i, cell in enumerate(row.cells):
                    text = cell.text.strip()
                    if text in replacements:
                        # 标签单元格找到后，取其右侧相邻的单元格作为值单元格
                        if i + 1 < len(row.cells):
                            target_cell = row.cells[i+1]

                            # 安全检查：如果右侧单元格也是标签，则跳过
                            # （避免"姓名|性别"这种紧邻标签的布局误填）
                            if target_cell.text.strip() in replacements:
                                continue

                            new_val = replacements[text]

                            # 通过修改第一个 run 的文本来保留原有字体格式
                            # （直接设置 cell.text 会丢失格式信息）
                            if target_cell.paragraphs:
                                p = target_cell.paragraphs[0]
                                if p.runs:
                                    p.runs[0].text = new_val
                                    # 清除多余的 run（可能是模板中的占位文本）
                                    for r in p.runs[1:]:
                                        r.text = ''
                                else:
                                    # 无 run 时创建新的
                                    p.add_run(new_val)

                                # 清除多余的段落
                                for p in target_cell.paragraphs[1:]:
                                    p.clear()
                            else:
                                target_cell.text = new_val

        # ---- 第二步：将学员证件照插入文档中的照片区域 ----
        if photo_path and os.path.exists(photo_path):
            _insert_photo_into_doc(doc, photo_path)

        doc.save(output_path)
        current_app.logger.info(f'Document generated: {output_path}')

    except Exception as e:
        current_app.logger.error(f'Failed to generate document: {str(e)}')
        raise


def _insert_photo_into_doc(doc, photo_path):
    """
    将照片插入到 Word 文档中。

    参数:
        doc: Document 对象
        photo_path: 照片文件路径
    """
    try:
        # 尝试将照片背景替换为白色（体检表要求白底证件照）
        # 如果背景替换失败（依赖库未安装等），降级使用原始照片
        temp_photo_path = None
        try:
            temp_fd, temp_photo_path = tempfile.mkstemp(suffix='.jpg')
            os.close(temp_fd)
            processed_photo_path = change_id_photo_bg(photo_path, temp_photo_path)
            photo_to_use = processed_photo_path if processed_photo_path == temp_photo_path else photo_path
        except Exception as e:
            current_app.logger.warning(f'Background processing failed, using original: {str(e)}')
            photo_to_use = photo_path
            temp_photo_path = None

        # 加载照片
        with open(photo_to_use, 'rb') as f:
            img_bytes = f.read()

        # 一寸证件照的标准尺寸（2.5cm × 3.5cm），转换为英寸
        default_w_in = 2.5 / 2.54  # 约 0.984 英寸
        default_h_in = 3.5 / 2.54  # 约 1.378 英寸

        # 在文档表格中查找照片插入位置
        replaced = False      # 标记是否已成功插入
        target_cell = None    # 目标单元格
        found_table = None    # 目标单元格所在的表格
        found_col_idx = None  # 目标单元格的列索引

        # 策略一：查找文本包含"照片"或"照"的单元格
        for ti, table in enumerate(doc.tables):
            for ri, row in enumerate(table.rows):
                for ci, cell in enumerate(row.cells):
                    txt = cell.text.strip()
                    if '照片' in txt or '照' in txt:
                        target_cell = cell
                        found_table = table
                        found_col_idx = ci
                        break
                if target_cell:
                    break
            if target_cell:
                break

        # 策略二（备选）：使用第一个表格的右上角单元格
        # 大多数体检表模板的照片区域位于右上角
        if not target_cell and len(doc.tables) > 0:
            tbl = doc.tables[0]
            target_cell = tbl.cell(0, len(tbl.rows[0].cells)-1)
            found_table = tbl
            found_col_idx = len(tbl.rows[0].cells)-1

        if target_cell:
            # 通过解析单元格的 XML 确定目标区域的实际宽度
            # Word 内部使用 OOXML 格式，宽度单位为 twip（1英寸=1440 twip）
            ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
            try:
                cell_root = etree.fromstring(target_cell._tc.xml.encode('utf-8'))
            except Exception:
                cell_root = None

            tcW = None
            if cell_root is not None:
                tcW = cell_root.xpath('.//w:tcPr/w:tcW', namespaces=ns)

            target_w_in = default_w_in
            if tcW:
                try:
                    w_val = int(tcW[0].get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}w'))
                    w_type = tcW[0].get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}type')
                    if w_type == 'dxa':
                        target_w_in = w_val / 1440.0
                except Exception:
                    pass
            else:
                # 备选：使用 gridCol 宽度
                try:
                    tbl_root = etree.fromstring(found_table._tbl.xml.encode('utf-8'))
                    gridCols = tbl_root.xpath('.//w:tblGrid/w:gridCol', namespaces=ns)
                    if gridCols and found_col_idx is not None and found_col_idx < len(gridCols):
                        gv = int(gridCols[found_col_idx].get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}w'))
                        target_w_in = gv / 1440.0
                except Exception:
                    pass

            # 校验目标宽度
            if not target_w_in or target_w_in <= 0 or target_w_in > 20:
                target_w_in = default_w_in

            tgt_w = target_w_in
            tgt_h = tgt_w * (3.5 / 2.5)  # 保持一寸证件照的标准纵横比 (2.5:3.5)

            # 确保最小尺寸
            min_w_in = 0.6
            min_h_in = 0.8
            if tgt_w < min_w_in:
                tgt_w = min_w_in
            if tgt_h < min_h_in:
                tgt_h = min_h_in

            # 调整尺寸并裁剪图片
            try:
                src_img = Image.open(io.BytesIO(img_bytes)).convert('RGB')
                src_w, src_h = src_img.size

                # 按 300 DPI 计算目标像素尺寸，确保打印质量
                dpi = 300
                px_w = max(120, int(tgt_w * dpi))
                px_h = max(160, int(tgt_h * dpi))
                # 等比缩放：取宽高方向较小的缩放比，保证图片完全显示
                scale = min(px_w / src_w, px_h / src_h)
                new_w = max(1, int(src_w * scale))
                new_h = max(1, int(src_h * scale))
                resized_img = src_img.resize((new_w, new_h), Image.LANCZOS)

                # 创建白色背景画布，将缩放后的图片居中粘贴
                # 这样即使原始图片比例不匹配，也会留白而非拉伸变形
                canvas = Image.new('RGB', (px_w, px_h), (255, 255, 255))
                paste_left = (px_w - new_w) // 2
                paste_top = (px_h - new_h) // 2
                canvas.paste(resized_img, (paste_left, paste_top))

                out_bio = io.BytesIO()
                canvas.save(out_bio, format='JPEG', quality=95)
                img_bytes_final = out_bio.getvalue()

                # 清除目标单元格中的原有内容（如占位文字"照片"）
                for p in list(target_cell.paragraphs):
                    p.clear()
                # 在单元格中居中插入处理后的图片
                p = target_cell.paragraphs[0]
                try:
                    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                except Exception:
                    pass
                run = p.add_run()
                bio2 = io.BytesIO(img_bytes_final)
                try:
                    run.add_picture(bio2, width=Inches(tgt_w), height=Inches(tgt_h))
                except TypeError:
                    # 兼容旧版 python-docx 不支持 height 参数的情况
                    bio2.seek(0)
                    run.add_picture(bio2, width=Inches(tgt_w))
                replaced = True
            except Exception as e:
                current_app.logger.error(f'Error preparing/resizing picture: {str(e)}')

        # 备选方案：如果无法找到照片单元格，尝试替换文档中已有的嵌入图片
        # 这种方式直接替换图片资源的二进制数据，不改变布局
        if not replaced:
            blip_rid = None
            # 方法一：从文档关系中查找图片类型的关系 ID
            for rel in doc.part.rels:
                try:
                    rel_obj = doc.part.rels[rel]
                    if getattr(rel_obj, 'reltype', '').endswith('/image') or \
                       'image' in getattr(rel_obj.target_part, '__class__', '').lower():
                        blip_rid = rel
                        break
                except Exception:
                    continue

            # 方法二：从段落 XML 中搜索嵌入图片的引用 ID
            if not blip_rid:
                import re
                for p in doc.paragraphs:
                    xml = p._element.xml
                    if 'r:embed="' in xml:
                        match = re.search(r'r:embed="([^\"]+)"', xml)
                        if match:
                            blip_rid = match.group(1)
                            break

            # 直接替换图片资源的二进制数据
            if blip_rid and blip_rid in doc.part.rels:
                try:
                    doc.part.rels[blip_rid].target_part._blob = img_bytes
                except Exception:
                    pass

    except Exception as e:
        current_app.logger.error(f'Error processing/replacing image: {str(e)}')
    finally:
        # 清理临时文件
        if temp_photo_path and os.path.exists(temp_photo_path):
            try:
                os.remove(temp_photo_path)
            except Exception:
                pass
